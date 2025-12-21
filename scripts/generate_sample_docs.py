"""
Script to generate sample PDF and Word documents for the knowledge base
"""
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_onboarding_pdf():
    """Create New Employee Onboarding Guide PDF"""
    output_path = "/home/claude/enterprise-knowledge-assistant/data/sample_docs/New_Employee_Onboarding.pdf"
    
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    Story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#002D72',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='#002D72',
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    title = Paragraph("Johns Hopkins University<br/>New Employee Onboarding Guide", title_style)
    Story.append(title)
    Story.append(Spacer(1, 0.3*inch))
    
    # Welcome
    Story.append(Paragraph("Welcome to Johns Hopkins!", heading_style))
    welcome_text = """
    Welcome to Johns Hopkins University! We're excited to have you join our community of 
    exceptional faculty, staff, and researchers. This guide will help you navigate your first 
    weeks and ensure a smooth transition into your new role.
    """
    Story.append(Paragraph(welcome_text, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # First Day
    Story.append(Paragraph("Your First Day", heading_style))
    first_day = """
    <b>What to Bring:</b><br/>
    • Government-issued photo ID (driver's license or passport)<br/>
    • Social Security card or equivalent documentation<br/>
    • Completed I-9 and tax forms (if not submitted electronically)<br/>
    • Bank account information for direct deposit setup<br/><br/>
    
    <b>What to Expect:</b><br/>
    Your first day will begin at 9:00 AM in the Human Resources Office (Garland Hall, 3rd Floor). 
    Plan to spend approximately 3-4 hours completing orientation activities, including:<br/>
    • HR paperwork and benefit elections<br/>
    • IT account setup and equipment distribution<br/>
    • Campus tour<br/>
    • ID badge creation<br/>
    • Meeting with your supervisor and team
    """
    Story.append(Paragraph(first_day, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # First Week
    Story.append(Paragraph("Your First Week Checklist", heading_style))
    checklist = """
    <b>Day 1-2:</b><br/>
    □ Complete HR orientation and paperwork<br/>
    □ Set up computer and accounts (email, VPN, Workday)<br/>
    □ Review department handbook and policies<br/>
    □ Schedule one-on-ones with key team members<br/>
    □ Familiarize yourself with building and facilities<br/><br/>
    
    <b>Day 3-5:</b><br/>
    □ Attend required training sessions (scheduled by HR)<br/>
    □ Review role expectations and goals with supervisor<br/>
    □ Join relevant Slack channels and distribution lists<br/>
    □ Explore campus resources and amenities<br/>
    □ Complete Information Security training (required within first week)<br/>
    □ Set up voicemail and email signature
    """
    Story.append(Paragraph(checklist, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # First 30 Days
    Story.append(Paragraph("Your First 30 Days", heading_style))
    thirty_days = """
    <b>Week 2-4 Goals:</b><br/>
    • Complete all required compliance training modules<br/>
    • Enroll in benefits (deadline: 31 days from hire date)<br/>
    • Attend new employee orientation session<br/>
    • Shadow colleagues in related roles<br/>
    • Begin working on initial projects<br/>
    • Schedule 30-day check-in with supervisor<br/>
    • Connect with your onboarding buddy (assigned by HR)<br/><br/>
    
    Your supervisor will provide a detailed 30-60-90 day plan with specific role objectives 
    and success metrics.
    """
    Story.append(Paragraph(thirty_days, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # IT Setup
    Story.append(Paragraph("Technology Setup", heading_style))
    it_setup = """
    <b>JHED ID:</b> Your Johns Hopkins Enterprise Directory (JHED) ID is your unique identifier 
    for all university systems. You'll receive this during orientation.<br/><br/>
    
    <b>Essential Systems:</b><br/>
    • <b>Email:</b> firstname.lastname@jhu.edu (Microsoft Outlook/Office 365)<br/>
    • <b>Workday:</b> HR system for payroll, time off, and personal information<br/>
    • <b>VPN:</b> Required for remote access to university resources<br/>
    • <b>Slack:</b> Primary communication tool for many departments<br/>
    • <b>Zoom:</b> Video conferencing (enterprise account)<br/>
    • <b>OneDrive/Box:</b> Cloud storage for university documents<br/><br/>
    
    <b>Password Requirements:</b><br/>
    • Minimum 12 characters<br/>
    • Must include uppercase, lowercase, number, and special character<br/>
    • Changes required every 90 days<br/>
    • Cannot reuse last 6 passwords<br/><br/>
    
    <b>IT Support:</b> For technical assistance, contact the Help Desk at 410-516-4357 
    or itsupport@jhu.edu.
    """
    Story.append(Paragraph(it_setup, styles['Normal']))
    
    Story.append(PageBreak())
    
    # Benefits
    Story.append(Paragraph("Benefits Overview", heading_style))
    benefits = """
    You must enroll in benefits within 31 days of your start date. If you miss this deadline, 
    you'll have to wait until the next open enrollment period (November).<br/><br/>
    
    <b>Key Benefits:</b><br/>
    • <b>Health Insurance:</b> Three plan options (PPO, HDHP, HMO)<br/>
    • <b>Dental & Vision:</b> Optional coverage through Delta Dental and VSP<br/>
    • <b>Retirement:</b> 403(b) with 5% employer match (immediate eligibility)<br/>
    • <b>Paid Time Off:</b> 15-25 vacation days plus 12 sick days annually<br/>
    • <b>Tuition Remission:</b> Available after 1 year of service<br/>
    • <b>Wellness Programs:</b> Free fitness center access and wellness incentives<br/><br/>
    
    <b>Benefits Enrollment:</b><br/>
    • Access the benefits portal at mybenefits.jhu.edu<br/>
    • Attend benefits orientation session (scheduled during first week)<br/>
    • Contact Benefits Office with questions: 410-516-2800<br/><br/>
    
    <b>Important:</b> Decisions made during initial enrollment are locked in for the calendar 
    year unless you experience a qualifying life event (marriage, birth, loss of coverage).
    """
    Story.append(Paragraph(benefits, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # Campus Resources
    Story.append(Paragraph("Campus Resources", heading_style))
    resources = """
    <b>Dining:</b> Multiple cafeterias, coffee shops, and restaurants across campus. 
    Use the JHU Mobile app to view menus and hours.<br/><br/>
    
    <b>Libraries:</b> Sheridan Libraries system offers research support, study spaces, 
    and extensive digital resources. Your JHED ID provides access.<br/><br/>
    
    <b>Recreation:</b> Free access to Recreation Center with fitness equipment, pools, 
    courts, and group classes. Towel service included.<br/><br/>
    
    <b>Parking & Transportation:</b><br/>
    • Parking permits: parking.jhu.edu (subsidized rates)<br/>
    • Shuttle service: Free between campuses<br/>
    • Transit subsidy: Up to $150/month pre-tax<br/>
    • Bike share: 50% discount on annual membership<br/><br/>
    
    <b>Safety:</b><br/>
    • Campus Security: 410-516-7777 (24/7)<br/>
    • Emergency: 911 or x-77777 from campus phone<br/>
    • NightRide shuttle service: Free evening transportation<br/>
    • Security escorts: Available upon request
    """
    Story.append(Paragraph(resources, styles['Normal']))
    Story.append(Spacer(1, 0.2*inch))
    
    # Important Contacts
    Story.append(Paragraph("Important Contacts", heading_style))
    contacts = """
    <b>Human Resources:</b> 410-516-2800 | hr@jhu.edu<br/>
    <b>IT Help Desk:</b> 410-516-4357 | itsupport@jhu.edu<br/>
    <b>Benefits Office:</b> 410-516-2800 | benefits@jhu.edu<br/>
    <b>Payroll:</b> 410-516-2830 | payroll@jhu.edu<br/>
    <b>Campus Security:</b> 410-516-7777<br/>
    <b>Employee Assistance Program:</b> 1-800-327-2251 (24/7)<br/><br/>
    
    <b>Online Resources:</b><br/>
    • HR Portal: hr.jhu.edu<br/>
    • IT Services: it.jhu.edu<br/>
    • Learning & Development: learning.jhu.edu<br/>
    • Campus Map: map.jhu.edu
    """
    Story.append(Paragraph(contacts, styles['Normal']))
    Story.append(Spacer(1, 0.3*inch))
    
    # Footer
    footer_text = """
    <i>Questions? Contact your HR representative or the New Employee Hotline at 410-516-2888.</i><br/>
    <i>Welcome to the Johns Hopkins family!</i>
    """
    Story.append(Paragraph(footer_text, styles['Normal']))
    
    # Build PDF
    doc.build(Story)
    print(f"✓ Created: {output_path}")

def create_security_guidelines_docx():
    """Create Data Security Guidelines Word document"""
    output_path = "/home/claude/enterprise-knowledge-assistant/data/sample_docs/Data_Security_Guidelines.docx"
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Johns Hopkins University', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 45, 114)
    
    subtitle = doc.add_heading('Data Security Guidelines', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(0, 45, 114)
    
    doc.add_paragraph('Version 3.2 | Effective November 2025', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Introduction
    doc.add_heading('Purpose', 2)
    doc.add_paragraph(
        'This document establishes security standards for protecting Johns Hopkins University '
        'data and information systems. All faculty, staff, students, and affiliates with access '
        'to university systems must comply with these guidelines.'
    )
    doc.add_paragraph()
    
    # Data Classification
    doc.add_heading('Data Classification', 2)
    doc.add_paragraph(
        'All JHU data must be classified into one of four categories based on sensitivity:'
    )
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Classification'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Examples'
    
    # Data rows
    classifications = [
        ['Public', 'Information intended for public disclosure', 'Published research, marketing materials, public websites'],
        ['Internal', 'Information for JHU community use only', 'Internal memos, policies, phone directories'],
        ['Confidential', 'Sensitive information requiring protection', 'Employee records, donor information, research data'],
        ['Restricted', 'Highly sensitive regulated information', 'SSN, health records (HIPAA), financial account numbers, research participant data']
    ]
    
    for i, row_data in enumerate(classifications, start=1):
        cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            cells[j].text = cell_text
    
    doc.add_paragraph()
    
    # Password Requirements
    doc.add_heading('Password Security', 2)
    doc.add_paragraph('Strong passwords are your first line of defense. JHU password requirements:')
    
    password_reqs = [
        'Minimum 12 characters in length',
        'Must contain uppercase letters (A-Z)',
        'Must contain lowercase letters (a-z)',
        'Must contain numbers (0-9)',
        'Must contain special characters (!@#$%^&*)',
        'Cannot contain your name or JHED ID',
        'Cannot reuse any of your last 6 passwords',
        'Must be changed every 90 days',
        'Account locks after 5 failed login attempts'
    ]
    
    for req in password_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Password Managers: Use of approved password managers (LastPass Enterprise, 1Password) '
        'is encouraged for managing complex passwords. Contact IT for access.'
    )
    doc.add_paragraph()
    
    # Multi-Factor Authentication
    doc.add_heading('Multi-Factor Authentication (MFA)', 2)
    doc.add_paragraph(
        'MFA is REQUIRED for all JHU accounts. MFA adds an extra layer of security by requiring '
        'two forms of verification: something you know (password) and something you have (phone or token).'
    )
    
    doc.add_paragraph('Approved MFA Methods:', style='List Bullet')
    mfa_methods = [
        'Duo Mobile app (recommended)',
        'SMS text message to registered phone',
        'Phone call to registered number',
        'Hardware security key (YubiKey)'
    ]
    for method in mfa_methods:
        p = doc.add_paragraph(method, style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Setup Instructions: Visit mfa.jhu.edu to configure your MFA method. '
        'Contact IT Support if you need assistance.'
    )
    doc.add_paragraph()
    
    # Device Security
    doc.add_heading('Device Security', 2)
    
    doc.add_heading('University-Owned Devices', 3)
    university_device_rules = [
        'Full disk encryption is mandatory on all laptops and mobile devices',
        'Automatic screen lock after 10 minutes of inactivity',
        'Keep operating system and software up to date with latest patches',
        'Install only approved software from JHU Software Center',
        'Enable "Find My Device" features for mobile devices',
        'Do not share devices with family members or non-JHU individuals',
        'Report lost or stolen devices immediately to IT Security'
    ]
    for rule in university_device_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Personal Devices (BYOD)', 3)
    doc.add_paragraph(
        'If you access JHU data on personal devices, you must comply with these requirements:'
    )
    personal_device_rules = [
        'Use only approved secure container apps (Outlook mobile, OneDrive)',
        'Enable device PIN/biometric lock',
        'Allow remote wipe capability if device is lost',
        'Do not download Confidential or Restricted data to personal devices',
        'Use JHU VPN when accessing university resources',
        'Understand that JHU may enforce security policies on your device'
    ]
    for rule in personal_device_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    
    # Network Security
    doc.add_heading('Network and Remote Access', 2)
    
    doc.add_paragraph('VPN Usage:')
    vpn_rules = [
        'VPN is REQUIRED when accessing JHU systems from off-campus',
        'Connect to VPN before accessing email, file shares, or databases',
        'Download Cisco AnyConnect from vpn.jhu.edu',
        'Use only JHU-approved VPN (do not use personal/commercial VPN services)',
        'Keep VPN client updated to latest version'
    ]
    for rule in vpn_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Wi-Fi Security:')
    wifi_rules = [
        'Use "JHU-Secure" network when on campus (not "JHU-Guest")',
        'Never use public Wi-Fi without VPN for university work',
        'Avoid "free Wi-Fi" at airports, cafes, hotels when handling sensitive data',
        'Disable automatic Wi-Fi connection on mobile devices'
    ]
    for rule in wifi_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    
    # Data Storage
    doc.add_heading('Data Storage and Sharing', 2)
    
    doc.add_paragraph('Approved Cloud Storage:')
    cloud_rules = [
        'OneDrive for Business: For individual and department files',
        'JHU Box: For collaboration and external sharing',
        'SharePoint: For team sites and document management',
        'Research Data Storage: For active research projects'
    ]
    for rule in cloud_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'PROHIBITED: Personal cloud services (Dropbox, Google Drive, iCloud) are NOT approved '
        'for JHU data. Use of unauthorized services may result in disciplinary action.',
        style='Intense Quote'
    )
    doc.add_paragraph()
    
    doc.add_paragraph('File Sharing Guidelines:')
    sharing_rules = [
        'Use JHU Box for sharing large files externally',
        'Set expiration dates on shared links (maximum 90 days)',
        'Require password protection for Confidential/Restricted data',
        'Never email Restricted data (SSN, financial accounts, health info)',
        'Use encrypted email (secure.jhu.edu) for Confidential information',
        'Verify recipient before sharing sensitive data'
    ]
    for rule in sharing_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_paragraph()
    
    # Phishing
    doc.add_heading('Phishing and Social Engineering', 2)
    doc.add_paragraph(
        'Phishing attacks are the #1 security threat. Be vigilant and report suspicious emails.'
    )
    
    doc.add_paragraph('Warning Signs of Phishing:')
    phishing_signs = [
        'Urgent or threatening language ("verify now or account will close")',
        'Generic greetings ("Dear User" instead of your name)',
        'Spelling and grammar errors',
        'Suspicious links (hover to preview - does domain look right?)',
        'Requests for passwords, SSN, or financial information',
        'Unexpected attachments or links',
        'Sender email doesn\'t match organization (paypal-secure.biz vs paypal.com)',
        'Offers that seem too good to be true'
    ]
    for sign in phishing_signs:
        doc.add_paragraph(sign, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How to Report Phishing: Forward suspicious emails to phishing@jhu.edu or click '
        'the "Report Phishing" button in Outlook. Do not click links or download attachments first!'
    )
    doc.add_paragraph()
    
    # Incident Response
    doc.add_heading('Security Incident Reporting', 2)
    doc.add_paragraph(
        'Report security incidents immediately - quick response can minimize damage.'
    )
    
    doc.add_paragraph('Reportable Incidents Include:')
    incidents = [
        'Lost or stolen devices (laptop, phone, tablet, USB drive)',
        'Suspected malware or virus infection',
        'Unauthorized access to your account',
        'Accidental data disclosure or breach',
        'Suspected phishing or social engineering',
        'Compromised passwords',
        'Physical security breaches',
        'Any suspicious activity involving JHU data or systems'
    ]
    for incident in incidents:
        doc.add_paragraph(incident, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Report To: Email security@jhu.edu or call 410-516-8888 (24/7 security hotline)'
    ).bold = True
    doc.add_paragraph()
    
    # Compliance
    doc.add_heading('Compliance and Training', 2)
    doc.add_paragraph('Required Training:')
    training = [
        'Information Security Awareness: Annual requirement for all users',
        'HIPAA Privacy & Security: For healthcare and research staff handling health information',
        'Data Privacy Training: For those handling personally identifiable information (PII)',
        'Role-specific training: Additional training based on data access level'
    ]
    for item in training:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Access training modules at learning.jhu.edu. Failure to complete required training '
        'may result in access suspension.'
    )
    doc.add_paragraph()
    
    doc.add_heading('Compliance Violations', 3)
    doc.add_paragraph(
        'Violations of security policies may result in:'
    )
    violations = [
        'Suspension of system access',
        'Disciplinary action up to and including termination',
        'Legal action if violations involve criminal activity',
        'Personal liability for damages resulting from negligence'
    ]
    for violation in violations:
        doc.add_paragraph(violation, style='List Bullet')
    
    doc.add_paragraph()
    
    # Contact Info
    doc.add_heading('Resources and Contact Information', 2)
    
    contacts_table = doc.add_table(rows=6, cols=2)
    contacts_table.style = 'Light List Accent 1'
    
    contacts_data = [
        ['IT Security Office', 'Email: security@jhu.edu | Phone: 410-516-8888'],
        ['IT Help Desk', 'Email: itsupport@jhu.edu | Phone: 410-516-4357'],
        ['Security Incident Hotline', '410-516-8888 (24/7)'],
        ['Phishing Reports', 'phishing@jhu.edu'],
        ['Security Training Portal', 'learning.jhu.edu/security'],
        ['Security Policies', 'policies.jhu.edu/security']
    ]
    
    for i, (label, info) in enumerate(contacts_data):
        cells = contacts_table.rows[i].cells
        cells[0].text = label
        cells[1].text = info
        # Make first column bold
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph(
        'This document is reviewed and updated quarterly. Last updated: November 2025',
        style='Normal'
    ).italic = True
    
    doc.add_paragraph(
        'For the complete Information Security Policy, visit: policies.jhu.edu/security',
        style='Normal'
    ).italic = True
    
    # Save document
    doc.save(output_path)
    print(f"✓ Created: {output_path}")

if __name__ == "__main__":
    print("Generating sample documents...")
    print()
    create_onboarding_pdf()
    create_security_guidelines_docx()
    print()
    print("✓ All sample documents created successfully!")
